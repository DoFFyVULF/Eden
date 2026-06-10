from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUT = "/Users/lukakataev/Desktop/REA/Eden/2.6_Оценка_экономической_эффективности_исправлено.docx"


def set_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc, text="", bold=False, center=False, first_line=1.25):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.first_line_indent = Pt(first_line * 28.3465) if first_line else None
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.line_spacing = 1.5
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, 12, bold=bold)
    return p


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    fmt.line_spacing = 1.5
    r = p.add_run(text)
    set_font(r, 12)
    return p


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(6)
    fmt.line_spacing = 1.5
    r = p.add_run(text)
    set_font(r, 12)
    return p


doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Pt(56.7)
    section.bottom_margin = Pt(56.7)
    section.left_margin = Pt(85.05)
    section.right_margin = Pt(42.525)

add_paragraph(doc, "2.6. Оценка экономической эффективности", bold=True, center=True, first_line=0)

add_paragraph(
    doc,
    "Экономическая оценка разработанной информационной системы выполнена с "
    "использованием показателей годового экономического эффекта, рентабельности "
    "инвестиций и срока окупаемости. Такой подход позволяет учитывать как разовые "
    "затраты на разработку и внедрение, так и ежегодные расходы на эксплуатацию "
    "и сопровождение системы."
)
add_paragraph(
    doc,
    "При оценке учтена специфика предприятий индустрии красоты. Переход на "
    "онлайн-запись позволяет принимать заявки вне рабочего времени, снижать "
    "нагрузку на администратора, ускорять формирование отчетности и отказаться "
    "от использования сторонних платных сервисов."
)
add_paragraph(
    doc,
    "Оценка проводится за годовой период, поскольку именно такой интервал "
    "позволяет наглядно определить срок окупаемости и практическую пользу от "
    "внедрения проекта. В расчетах приняты следующие допущения: средняя стоимость "
    "подписки на внешний сервис онлайн-записи составляет 5 000 руб. в месяц, "
    "дополнительный поток клиентов благодаря онлайн-записи составляет 5 человек "
    "в месяц, администратор экономит в среднем 1,5 часа рабочего времени в день "
    "на обработке записей клиентов, а также 6 часов в месяц на формировании отчетности."
)

add_paragraph(doc, "2.6.1. Затраты на разработку и внедрение", bold=True, first_line=0)
add_paragraph(
    doc,
    "К единовременным затратам относятся расходы на оплату труда разработчиков, "
    "страховые взносы и накладные расходы, связанные с реализацией проекта. "
    "Основной объем вложений формируется на этапе создания и подготовки системы "
    "к эксплуатации."
)
add_paragraph(doc, "Трудозатраты на разработку рассчитываются по формуле:")
add_formula(doc, "Зот = Σ(Ti × Cч,i),")
add_paragraph(
    doc,
    "где Ti - трудоемкость i-го этапа работ, чел.-ч; Cч,i - часовая ставка "
    "специалиста, руб./ч."
)
add_paragraph(doc, "В таблице 45 представлены затраты на разработку системы.")

add_table_caption(doc, "Таблица 45 - Единоразовые затраты на разработку системы")
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Тип затрат"
hdr[1].text = "Объем работы"
hdr[2].text = "Ставка, руб./ч"
hdr[3].text = "Итог, руб."
rows = [
    ["Анализ требований и проектирование", "24 ч", "500", "12 000"],
    ["Разработка серверной части", "80 ч", "500", "40 000"],
    ["Разработка клиентской части", "96 ч", "500", "48 000"],
    ["Тестирование системы", "36 ч", "450", "16 200"],
    ["Страховые взносы", "30 % от ФОТ", "-", "34 860"],
    ["Накладные расходы", "20 % от ФОТ", "-", "23 240"],
    ["ИТОГО", "", "", "174 300"],
]
for row in rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value

add_paragraph(
    doc,
    "Таким образом, для разработки и внедрения информационной системы требуется "
    "174 300 руб. единовременных вложений."
)

add_paragraph(doc, "2.6.2. Экономический эффект от внедрения", bold=True, first_line=0)
add_paragraph(
    doc,
    "Экономический эффект от внедрения системы формируется за счет четырех "
    "основных составляющих: отказа от использования платного аналога для "
    "онлайн-записи, получения дополнительной прибыли благодаря новым клиентам, "
    "сокращения трудозатрат администратора на обработку записей и сокращения "
    "времени на формирование отчетности."
)
add_paragraph(
    doc,
    "Годовая экономия на отказе от внешнего SaaS-сервиса определяется следующим образом:"
)
add_formula(doc, "Э1 = 5 000 × 12 = 60 000 руб./год.")
add_paragraph(
    doc,
    "Дополнительная прибыль от роста числа записей, поступающих в нерабочее время, "
    "рассчитывается по формуле:"
)
add_formula(doc, "Э2 = 5 × 2 500 × 0,6 × 12 = 90 000 руб./год.")
add_paragraph(
    doc,
    "Экономия трудозатрат администратора на обработке записей клиентов составляет:"
)
add_formula(doc, "Э3 = 1,5 × 22 × 12 × 300 = 118 800 руб./год.")
add_paragraph(
    doc,
    "Экономия времени на формировании отчетности определяется следующим образом:"
)
add_formula(doc, "Э4 = 6 × 12 × 300 = 21 600 руб./год.")
add_paragraph(doc, "Общий годовой экономический эффект определяется как сумма составляющих:")
add_formula(doc, "Эгод = Э1 + Э2 + Э3 + Э4 = 60 000 + 90 000 + 118 800 + 21 600 = 290 400 руб./год.")

add_paragraph(doc, "2.6.3. Эксплуатационные расходы", bold=True, first_line=0)
add_paragraph(
    doc,
    "Помимо единовременных затрат, в процессе использования системы возникают "
    "ежегодные расходы на ее поддержку и размещение. В качестве инфраструктурной "
    "платформы выбран провайдер Selectel [38]."
)
add_paragraph(doc, "Для функционирования информационной системы достаточно конфигурации, представленной в таблице 46.")

add_table_caption(doc, "Таблица 46 - Характеристика серверной конфигурации")
table2 = doc.add_table(rows=1, cols=2)
table2.style = "Table Grid"
hdr2 = table2.rows[0].cells
hdr2[0].text = "Параметр"
hdr2[1].text = "Значение"
rows2 = [
    ["Платформа", "Облачный сервер"],
    ["Тип конфигурации", "Standard Line"],
    ["Процессор (vCPU)", "1 × 2,30 ГГц"],
    ["Оперативная память (RAM)", "2 ГБ"],
    ["Дисковое пространство (SSD)", "16 ГБ"],
    ["Стоимость в год", "17 485 руб."],
]
for row in rows2:
    cells = table2.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value

add_paragraph(
    doc,
    "Данная конфигурация обеспечивает необходимый запас производительности с учетом "
    "пиковой нагрузки в вечерние часы и в выходные дни. Помимо этого, требуется "
    "обеспечить сервер доменным именем: по состоянию на апрель 2026 года стоимость "
    "регистрации домена «eden-beauty.ru» составляет 169 руб., а ежегодное продление "
    "обойдется примерно в 600-800 руб."
)
add_paragraph(doc, "Рассмотрим в таблице 47 годовые эксплуатационные расходы.")

add_table_caption(doc, "Таблица 47 - Годовые эксплуатационные расходы")
table3 = doc.add_table(rows=1, cols=3)
table3.style = "Table Grid"
hdr3 = table3.rows[0].cells
hdr3[0].text = "№"
hdr3[1].text = "Статья расходов"
hdr3[2].text = "Стоимость, руб./год"
rows3 = [
    ["1", "Аренда VPS-сервера", "17 485"],
    ["2", "Продление доменного имени", "800"],
    ["3", "Резервное копирование", "2 000"],
    ["4", "Техническая поддержка и администрирование", "10 000"],
    ["", "ИТОГО", "30 285"],
]
for row in rows3:
    cells = table3.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value

add_paragraph(doc, "Тогда чистый годовой экономический эффект будет равен:")
add_formula(doc, "Эчист = Эгод - Ээксп = 290 400 - 30 285 = 260 115 руб./год.")

add_paragraph(doc, "2.6.4. Расчет показателей эффективности", bold=True, first_line=0)
add_paragraph(
    doc,
    "Для оценки эффективности проекта используются два основных показателя: "
    "срок окупаемости и рентабельность инвестиций."
)
add_paragraph(doc, "Срок окупаемости определяется по формуле:")
add_formula(doc, "T = K / Эчист")
add_formula(doc, "T = 174 300 / 260 115 = 0,67 года.")
add_paragraph(
    doc,
    "Следовательно, срок окупаемости проекта составляет около 0,67 года, "
    "или примерно 8 месяцев."
)
add_paragraph(
    doc,
    "Показатель рентабельности инвестиций рассчитывается как отношение чистого "
    "годового экономического эффекта к первоначальным инвестициям:"
)
add_formula(doc, "ROI = (Эчист / K) × 100 %")
add_formula(doc, "ROI = (260 115 / 174 300) × 100 % = 149,23 %.")
add_paragraph(
    doc,
    "Полученные значения свидетельствуют о том, что внедрение разработанной "
    "системы является экономически целесообразным решением. Проект окупается "
    "менее чем за один год, а чистый годовой экономический эффект превышает "
    "величину первоначальных инвестиций."
)

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    set_font(r, 11)

doc.save(OUT)
print(OUT)
