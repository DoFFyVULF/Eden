from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm


OUT = "/Users/lukakataev/Desktop/REA/Eden/СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ_по_методичке.docx"


sources = {
    1: "Федеральный закон от 27.07.2006 № 149-ФЗ (ред. от 29.12.2025) «Об информации, информационных технологиях и о защите информации». Текст : электронный // КонсультантПлюс : [сайт]. URL: https://www.consultant.ru/document/cons_doc_LAW_61798/ (дата обращения: 03.05.2026).",
    2: "Федеральный закон от 27.07.2006 № 152-ФЗ (ред. от 29.12.2025) «О персональных данных». Текст : электронный // КонсультантПлюс : [сайт]. URL: https://www.consultant.ru/document/cons_doc_LAW_61801/ (дата обращения: 03.05.2026).",
    3: "Закон РФ от 07.02.1992 № 2300-1 (ред. от 28.12.2025) «О защите прав потребителей». Текст : электронный // КонсультантПлюс : [сайт]. URL: https://www.consultant.ru/document/cons_doc_LAW_305/ (дата обращения: 03.05.2026).",
    4: "ГОСТ Р 56939-2024. Защита информации. Разработка безопасного программного обеспечения. Общие требования. Текст : электронный // Росстандарт : [сайт]. URL: https://protect.gost.ru/gost/details/f3818925-a96f-4f55-96e9-46b44720ee64 (дата обращения: 03.05.2026).",
    5: "ГОСТ Р 59347-2021. Системная инженерия. Защита информации в процессе определения архитектуры системы. Текст : электронный // Росстандарт : [сайт]. URL: https://protect.gost.ru/gost/details/4c5fa9e1-380b-427a-8da1-2247c6b0eb7d (дата обращения: 03.05.2026).",
    6: "ГОСТ Р 59547-2021. Защита информации. Мониторинг информационной безопасности. Общие положения. Текст : электронный // Росстандарт : [сайт]. URL: https://protect.gost.ru/gost/details/c584277d-3b7d-4459-96bb-f6441a0ecc98 (дата обращения: 03.05.2026).",
    7: "ГОСТ Р ИСО/МЭК 27002-2021. Информационные технологии. Методы и средства обеспечения безопасности. Свод норм и правил применения мер обеспечения информационной безопасности. Текст : электронный // Росстандарт : [сайт]. URL: https://protect.gost.ru/gost/details/be00167c-6e26-4042-ac31-36be8a9d6b2e (дата обращения: 03.05.2026).",
    8: "NestJS Documentation. Текст : электронный // NestJS Docs : [сайт]. URL: https://docs.nestjs.com/ (дата обращения: 03.05.2026).",
    9: "NestJS Documentation: Controllers. Текст : электронный // NestJS Docs : [сайт]. URL: https://docs.nestjs.com/controllers (дата обращения: 03.05.2026).",
    10: "NestJS Documentation: Authentication. Текст : электронный // NestJS Docs : [сайт]. URL: https://docs.nestjs.com/security/authentication (дата обращения: 03.05.2026).",
    11: "React Documentation. Текст : электронный // React : [сайт]. URL: https://react.dev/ (дата обращения: 03.05.2026).",
    12: "The TypeScript Handbook. Текст : электронный // TypeScript : [сайт]. URL: https://www.typescriptlang.org/docs/handbook/intro.html (дата обращения: 03.05.2026).",
    13: "Next.js Documentation. Текст : электронный // Next.js : [сайт]. URL: https://nextjs.org/docs (дата обращения: 03.05.2026).",
    14: "PostgreSQL Documentation. Текст : электронный // PostgreSQL : [сайт]. URL: https://www.postgresql.org/docs/current/ (дата обращения: 03.05.2026).",
    15: "Prisma Documentation. Текст : электронный // Prisma : [сайт]. URL: https://www.prisma.io/docs (дата обращения: 03.05.2026).",
    16: "Prisma Schema. Текст : электронный // Prisma : [сайт]. URL: https://www.prisma.io/docs/orm/prisma-schema (дата обращения: 03.05.2026).",
    17: "Prisma Client. Текст : электронный // Prisma : [сайт]. URL: https://www.prisma.io/docs/orm/prisma-client (дата обращения: 03.05.2026).",
    18: "Node.js Documentation. Текст : электронный // Node.js : [сайт]. URL: https://nodejs.org/en/docs (дата обращения: 03.05.2026).",
    19: "Docker Overview. Текст : электронный // Docker Docs : [сайт]. URL: https://docs.docker.com/get-started/docker-overview/ (дата обращения: 03.05.2026).",
    20: "Docker Compose Documentation. Текст : электронный // Docker Docs : [сайт]. URL: https://docs.docker.com/compose/ (дата обращения: 03.05.2026).",
    21: "Axios Documentation. Текст : электронный // Axios : [сайт]. URL: https://axios-http.com/docs/intro (дата обращения: 03.05.2026).",
    22: "TanStack Query Documentation for React. Текст : электронный // TanStack : [сайт]. URL: https://tanstack.com/query/latest/docs/framework/react/overview (дата обращения: 03.05.2026).",
    23: "Tailwind CSS Documentation. Текст : электронный // Tailwind CSS : [сайт]. URL: https://tailwindcss.com/docs (дата обращения: 03.05.2026).",
    24: "React Hook Form Documentation. Текст : электронный // React Hook Form : [сайт]. URL: https://react-hook-form.com/docs (дата обращения: 03.05.2026).",
    25: "Recharts Documentation. Текст : электронный // Recharts : [сайт]. URL: https://recharts.github.io/ (дата обращения: 03.05.2026).",
    26: "Git Documentation. Текст : электронный // Git : [сайт]. URL: https://git-scm.com/docs (дата обращения: 03.05.2026).",
    27: "NGINX Beginner’s Guide. Текст : электронный // NGINX : [сайт]. URL: https://nginx.org/en/docs/beginners_guide.html (дата обращения: 03.05.2026).",
    28: "Certbot Instructions. Текст : электронный // Certbot : [сайт]. URL: https://certbot.eff.org/instructions (дата обращения: 03.05.2026).",
    29: "OWASP API Security Top 10. Текст : электронный // OWASP Foundation : [сайт]. URL: https://owasp.org/www-project-api-security/ (дата обращения: 03.05.2026).",
    30: "OWASP Top 10:2021. Текст : электронный // OWASP Foundation : [сайт]. URL: https://owasp.org/www-project-top-ten/ (дата обращения: 03.05.2026).",
    31: "YCLIENTS: онлайн-запись и автоматизация бизнеса услуг. Текст : электронный // YCLIENTS : [сайт]. URL: https://www.yclients.com/ (дата обращения: 03.05.2026).",
    32: "DIKIDI Business: онлайн-запись и управление клиентским сервисом. Текст : электронный // DIKIDI : [сайт]. URL: https://dikidi.ru/ru (дата обращения: 03.05.2026).",
    33: "1С:Предприятие 8. Салон красоты. Описание решения. Текст : электронный // Фирма «1С» : [сайт]. URL: https://solutions.1c.ru/catalog/beauty-salon (дата обращения: 03.05.2026).",
    34: "Cross-Origin Resource Sharing (CORS). Текст : электронный // MDN Web Docs : [сайт]. URL: https://developer.mozilla.org/en-US/docs/Web/HTTP/CORS (дата обращения: 03.05.2026).",
    35: "Fielding R., Nottingham M., Reschke J. HTTP Semantics. RFC 9110. Текст : электронный // IETF Datatracker : [сайт]. URL: https://datatracker.ietf.org/doc/html/rfc9110 (дата обращения: 03.05.2026).",
    36: "JSON Web Tokens. Текст : электронный // Auth0 Docs : [сайт]. URL: https://auth0.com/docs/tokens/concepts/jwts (дата обращения: 03.05.2026).",
    37: "A guide to IDEF diagrams. Текст : электронный // Lucidchart : [сайт]. URL: https://www.lucidchart.com/blog/idef-diagrams (дата обращения: 03.05.2026).",
    38: "Виртуальные сервера VDS. Текст : электронный // timeweb : [сайт]. URL: https://timeweb.com/ru/services/vds (дата обращения: 03.05.2026).",
}


new_order = [
    1, 2, 3, 4, 5, 6, 7,
    37, 21, 28, 34, 32, 20, 19, 35, 26, 36, 8, 10, 9, 13, 27, 18,
    29, 30, 14, 17, 15, 16, 11, 24, 25, 23, 22, 12, 31, 33, 38,
]


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(1.5)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(14)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(12)
run = title.add_run("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ И ЛИТЕРАТУРЫ")
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(14)

for new_no, old_no in enumerate(new_order, start=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{new_no}. {sources[old_no]}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)

doc.save(OUT)
print(OUT)
